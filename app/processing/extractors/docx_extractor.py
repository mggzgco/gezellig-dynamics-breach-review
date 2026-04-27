import io
from docx import Document
import logging

from app.processing.extractors.table_renderer import render_rows
from app.processing.extractors.types import ExtractedText

logger = logging.getLogger(__name__)


def extract_with_metadata(data: bytes) -> ExtractedText:
    """Extract DOCX text while preserving tables as labeled record rows."""
    try:
        doc = Document(io.BytesIO(data))
        parts = []
        table_count = 0

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # Extract tables
        for index, table in enumerate(doc.tables, start=1):
            table_count += 1
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            parts.extend(render_rows(f"Table {index}", table_rows))

        return ExtractedText(
            text="\n".join(parts),
            extraction_method="docx_text_and_tables",
            parser="python-docx",
            table_count=table_count,
            structured=table_count > 0,
        )
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ExtractedText(
            extraction_method="docx_text_and_tables",
            parser="python-docx",
            warnings=[str(e)[:200]],
        )


def extract(data: bytes) -> str:
    """Backward-compatible DOCX extraction API."""
    return extract_with_metadata(data).text
