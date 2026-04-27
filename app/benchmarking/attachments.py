from __future__ import annotations

import csv
import io
import textwrap
import zipfile
from email import policy
from email.message import EmailMessage
from email.mime.message import MIMEMessage
from email.parser import BytesParser
from typing import Optional

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont

from app.benchmarking.fixtures import (
    ATTACHMENT_ROTATION,
    ORG_NAMES,
    SUPPORT_EMAILS,
    SUPPORT_PHONES,
    TICKET_CODES,
    AttachmentSpec,
    FixturePerson,
    _field_lines,
    _field_pairs_for_attachment,
    _negative_lines,
)
from app.benchmarking.ground_truth import ExpectedFinding


def _compose_body(subject: str, greeting: str, lines: list[str], *, html_only: bool = False) -> EmailMessage:
    message = EmailMessage()
    message["Subject"] = subject
    message["From"] = "notifications@records.example"
    message["To"] = "reviewer@example.com"
    plain_lines = [greeting, "", *lines]
    plain_text = "\n".join(plain_lines)
    if html_only:
        html_body = "<html><body>" + "".join(f"<p>{line}</p>" for line in plain_lines if line) + "</body></html>"
        message.add_alternative(html_body, subtype="html")
    else:
        message.set_content(plain_text)
    return message


def _text_attachment(lines: list[str]) -> bytes:
    return ("\n".join(lines) + "\n").encode("utf-8")


def _csv_attachment(headers: list[str], rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(headers)
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def _xlsx_attachment(sheet_name: str, headers: list[str], rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = sheet_name
    worksheet.append(headers)
    for row in rows:
        worksheet.append(row)
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _docx_attachment(
    title: str,
    paragraphs: list[str],
    headers: Optional[list[str]] = None,
    rows: Optional[list[list[str]]] = None,
) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if headers and rows:
        table = document.add_table(rows=1, cols=len(headers))
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _rtf_attachment(lines: list[str]) -> bytes:
    escaped = [line.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}") for line in lines]
    body = r"\par ".join(escaped)
    return ("{\\rtf1\\ansi " + body + "}").encode("utf-8")


def _image_attachment_bytes(
    title: str,
    lines: list[str],
    *,
    image_format: str = "PNG",
    low_contrast: bool = False,
) -> bytes:
    image = _render_document_image(title, lines, low_contrast=low_contrast)
    output = io.BytesIO()
    save_kwargs = {"format": image_format.upper()}
    if image_format.upper() == "JPEG":
        image = image.convert("RGB")
        save_kwargs["quality"] = 92
    image.save(output, **save_kwargs)
    return output.getvalue()


def _pdf_attachment_bytes(title: str, pages: list[list[str]], *, low_contrast: bool = False) -> bytes:
    rendered_pages = [
        _render_document_image(
            title if page_index == 0 else f"{title} (continued)",
            page_lines,
            low_contrast=low_contrast,
        ).convert("RGB")
        for page_index, page_lines in enumerate(pages)
    ]
    if not rendered_pages:
        rendered_pages = [_render_document_image(title, ["No content provided."], low_contrast=low_contrast).convert("RGB")]
    output = io.BytesIO()
    rendered_pages[0].save(output, format="PDF", save_all=True, append_images=rendered_pages[1:])
    return output.getvalue()


def _zip_attachment(members: list[AttachmentSpec]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for member in members:
            archive.writestr(member.filename, member.data)
    return output.getvalue()


def _eml_bytes(
    subject: str,
    body_lines: list[str],
    attachments: Optional[list[AttachmentSpec]] = None,
    *,
    html_only: bool = False,
) -> bytes:
    message = _compose_body(subject, "Please review the attached information.", body_lines, html_only=html_only)
    if attachments:
        for attachment in attachments:
            _add_attachment(message, attachment)
    return message.as_bytes()


def _add_attachment(message: EmailMessage, attachment: AttachmentSpec) -> None:
    maintype, subtype = attachment.mime_type.split("/", 1)
    message.add_attachment(attachment.data, maintype=maintype, subtype=subtype, filename=attachment.filename)


def _add_nested_eml_attachment(message: EmailMessage, filename: str, data: bytes) -> None:
    nested_message = BytesParser(policy=policy.default).parsebytes(data)
    message.make_mixed()
    nested_part = MIMEMessage(nested_message)
    nested_part.add_header("Content-Disposition", "attachment", filename=filename)
    message.attach(nested_part)


def _set_single_header(message: EmailMessage, header_name: str, value: str) -> None:
    if header_name in message:
        message.replace_header(header_name, value)
    else:
        message[header_name] = value


def _expected_findings_from_pairs(
    source_ref: str,
    pairs: list[tuple[str, str] | tuple[str, str, str]],
    *,
    attachment_filename: Optional[str] = None,
    location_kind: str = "body",
    default_entity_name: Optional[str] = None,
) -> list[ExpectedFinding]:
    findings: list[ExpectedFinding] = []
    for pair in pairs:
        if len(pair) == 3:
            pii_type, raw_value, entity_name = pair
        else:
            pii_type, raw_value = pair
            entity_name = default_entity_name
        findings.append(
            ExpectedFinding(
                pii_type=pii_type,
                raw_value=raw_value,
                source_ref=source_ref,
                entity_name=entity_name,
                attachment_filename=attachment_filename,
                location_kind=location_kind,
            )
        )
    return findings


def _negative_attachment(index: int) -> AttachmentSpec:
    kind = ATTACHMENT_ROTATION[index % len(ATTACHMENT_ROTATION)]
    lines = _negative_lines(index) + [f"Document owner: {ORG_NAMES[(index + 2) % len(ORG_NAMES)]}"]
    filename = f"operations_{index:03d}.{kind}"
    if kind == "txt":
        return AttachmentSpec(filename=filename, mime_type="text/plain", data=_text_attachment(lines))
    if kind == "csv":
        headers = ["ticket", "owner", "contact", "review_date", "source_ip"]
        rows = [[
            TICKET_CODES[index % len(TICKET_CODES)],
            ORG_NAMES[index % len(ORG_NAMES)],
            SUPPORT_EMAILS[index % len(SUPPORT_EMAILS)],
            f"2026-0{(index % 9) + 1}-15",
            f"203.{index % 50}.10.24",
        ]]
        return AttachmentSpec(filename=filename, mime_type="text/csv", data=_csv_attachment(headers, rows))
    if kind == "docx":
        return AttachmentSpec(
            filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            data=_docx_attachment("Operations Update", lines),
        )
    if kind == "xlsx":
        headers = ["ticket", "team", "status", "review_date", "support_email"]
        rows = [[
            TICKET_CODES[index % len(TICKET_CODES)],
            ORG_NAMES[index % len(ORG_NAMES)],
            "Open",
            f"2026-0{(index % 9) + 1}-20",
            SUPPORT_EMAILS[index % len(SUPPORT_EMAILS)],
        ]]
        return AttachmentSpec(
            filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            data=_xlsx_attachment("Ops", headers, rows),
        )
    return AttachmentSpec(filename=filename, mime_type="text/rtf", data=_rtf_attachment(lines))


def _positive_attachment_content(
    index: int,
    person: FixturePerson,
    bundle: list[str],
    kind: str,
) -> tuple[AttachmentSpec, list[tuple[str, str]]]:
    field_lines, findings = _field_lines(person, bundle, index)
    field_pairs = _field_pairs_for_attachment(person, bundle, index)
    attachment_filename = f"record_{index:03d}.{kind if kind != 'nested' else 'eml'}"

    if kind == "txt":
        return (
            AttachmentSpec(
                filename=attachment_filename,
                mime_type="text/plain",
                data=_text_attachment(["Record Summary", *field_lines]),
            ),
            findings,
        )

    if kind == "csv":
        if index % 2 == 0:
            headers = [label for label, _ in field_pairs]
            rows = [[value for _, value in field_pairs]]
        else:
            headers = ["Field", "Value"]
            rows = [[label, value] for label, value in field_pairs]
        return (
            AttachmentSpec(filename=attachment_filename, mime_type="text/csv", data=_csv_attachment(headers, rows)),
            findings,
        )

    if kind == "xlsx":
        if index % 2 == 0:
            headers = [label.lower().replace(" ", "_") for label, _ in field_pairs]
            rows = [[value for _, value in field_pairs]]
        else:
            headers = ["field", "value"]
            rows = [[label, value] for label, value in field_pairs]
        return (
            AttachmentSpec(
                filename=attachment_filename,
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                data=_xlsx_attachment("Record", headers, rows),
            ),
            findings,
        )

    if kind == "docx":
        return (
            AttachmentSpec(
                filename=attachment_filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=_docx_attachment("Attached Record", ["Please review the attached information.", *field_lines]),
            ),
            findings,
        )

    if kind == "rtf":
        return (
            AttachmentSpec(
                filename=attachment_filename,
                mime_type="text/rtf",
                data=_rtf_attachment(["Attached Record", *field_lines]),
            ),
            findings,
        )

    if kind == "nested":
        nested_bytes = _eml_bytes(
            "Forwarded Record Details",
            ["Forwarding the enclosed record.", *field_lines],
            html_only=index % 3 == 0,
        )
        return (
            AttachmentSpec(filename=attachment_filename, mime_type="message/rfc822", data=nested_bytes),
            findings,
        )

    if kind == "zip":
        member_kind = ATTACHMENT_ROTATION[index % len(ATTACHMENT_ROTATION)]
        member_attachment, member_findings = _positive_attachment_content(index, person, bundle, member_kind)
        noise_member = AttachmentSpec(
            filename=f"ops_note_{index:03d}.txt",
            mime_type="text/plain",
            data=_text_attachment(_negative_lines(index)),
        )
        return (
            AttachmentSpec(
                filename=f"review_packet_{index:03d}.zip",
                mime_type="application/zip",
                data=_zip_attachment([member_attachment, noise_member]),
            ),
            member_findings,
        )

    raise ValueError(f"Unsupported positive attachment kind: {kind}")


def _render_document_image(title: str, lines: list[str], *, low_contrast: bool = False) -> Image.Image:
    """Render a simple form-like page image for OCR and scanned-document scenarios."""
    font = ImageFont.load_default()
    margin = 40
    content_width = 118
    rendered_lines: list[str] = []
    for line in lines:
        stripped = line.strip() or " "
        wrapped = textwrap.wrap(stripped, width=content_width) or [" "]
        rendered_lines.extend(wrapped)

    line_height = 24
    title_height = 30
    height = margin * 2 + title_height + max(1, len(rendered_lines)) * line_height + 24

    if low_contrast:
        background = (246, 244, 238)
        panel = (241, 238, 230)
        ink = (84, 84, 84)
        accent = (112, 112, 112)
    else:
        background = (255, 255, 255)
        panel = (248, 250, 252)
        ink = (22, 28, 36)
        accent = (76, 92, 120)

    image = Image.new("RGB", (1400, height), background)
    draw = ImageDraw.Draw(image)
    draw.rectangle((16, 16, 1384, height - 16), outline=accent, width=2, fill=panel)
    draw.text((margin, margin), title, fill=ink, font=font)
    draw.line((margin, margin + title_height, 1360, margin + title_height), fill=accent, width=2)

    y = margin + title_height + 16
    for line in rendered_lines:
        draw.text((margin, y), line, fill=ink, font=font)
        y += line_height
    return image
