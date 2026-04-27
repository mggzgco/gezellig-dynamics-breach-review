import tempfile
import unittest
import zipfile
from email.message import EmailMessage
from io import BytesIO
from pathlib import Path

import openpyxl
from docx import Document

from app.models import Attachment
from app.processing.attachment_handler import extract_text_from_attachment
from app.processing.extractors import xlsx_extractor
from app.processing.pii_engine import scan_text
from app.processing.pipeline import process_single_eml


class PIIDetectionQualityTests(unittest.TestCase):
    def test_quote_text_does_not_trigger_drivers_license(self):
        text = (
            "Thanks for your interest in our product line.\n"
            "Item: Enterprise API License (50 seats)\n"
            "Unit price: $2,500\n"
            "Quantity: 50\n"
            "Subtotal: $125,000\n"
            "Valid until: May 20, 2026\n"
        )

        matches = scan_text(text, "quote.txt")
        self.assertFalse(any(match.pii_type == "DRIVERS_LICENSE" for match in matches))

    def test_birth_context_is_word_based_not_substring_based(self):
        text = (
            "Congratulations on your upcoming retirement on May 31, 2026.\n"
            "Health coverage can continue under COBRA.\n"
        )

        matches = scan_text(text, "benefits.txt")
        self.assertFalse(any(match.pii_type == "DOB" for match in matches))

    def test_person_record_date_detects_dob_without_explicit_birth_label(self):
        text = (
            "Patient record follows:\n"
            "Jane Smith, 04/15/2026, 123 Main Street, Portland, OR 97205\n"
            "Phone: 212-555-0189\n"
        )

        matches = scan_text(text, "record.txt")
        dob_matches = [match for match in matches if match.pii_type == "DOB"]

        self.assertEqual(1, len(dob_matches))
        self.assertEqual("2026-04-15", dob_matches[0].normalized_value)

    def test_deadline_dates_do_not_trigger_dob(self):
        text = (
            "Important filing deadlines approaching:\n"
            "Form 10-K (annual): Due May 1, 2026\n"
            "Provide audit documentation by April 30, 2026.\n"
        )

        matches = scan_text(text, "deadline.txt")
        self.assertFalse(any(match.pii_type == "DOB" for match in matches))

    def test_birth_date_context_does_not_bleed_to_other_date_fields(self):
        text = (
            "Child: Emma Rodriguez\n"
            "Birth date: March 15, 2019 Adoption date: April 20, 2026\n"
        )

        matches = scan_text(text, "family.txt")
        dob_values = sorted(match.normalized_value for match in matches if match.pii_type == "DOB")

        self.assertEqual(["2019-03-15"], dob_values)

    def test_phone_detection_is_deduplicated(self):
        matches = scan_text("Phone: 1-212-555-1212", "contact.txt")
        phone_matches = [match for match in matches if match.pii_type == "PHONE"]

        self.assertEqual(1, len(phone_matches))
        self.assertEqual("2125551212", phone_matches[0].normalized_value)

    def test_service_contact_details_are_not_treated_as_leaked_contact_pii(self):
        text = (
            "Questions: facilities@university.edu or 503-555-0147\n"
            "Support: helpdesk@company.com\n"
            "Contact recall@manufacturer.com for refund processing.\n"
        )

        matches = scan_text(text, "service_notice.txt")
        detected_types = {match.pii_type for match in matches}

        self.assertFalse({"EMAIL", "PHONE"} & detected_types)

    def test_compact_account_number_is_not_mislabeled_as_ssn_or_bank_account(self):
        text = (
            "Member,\n"
            "Your auto loan payment is due.\n"
            "Account #: 334445829\n"
            "Due date: April 25, 2026\n"
        )

        matches = scan_text(text, "loan_notice.txt")
        detected_types = {match.pii_type for match in matches}

        self.assertNotIn("SSN", detected_types)
        self.assertNotIn("BANK_ACCOUNT", detected_types)

    def test_infrastructure_ip_alert_is_not_treated_as_pii(self):
        text = (
            "Network Team,\n"
            "Intrusion detection alert from external IP:\n"
            "Source IP: 203.45.67.89\n"
            "Port: 22 (SSH)\n"
        )

        matches = scan_text(text, "infra_alert.txt")
        self.assertFalse(any(match.pii_type == "IPV4" for match in matches))

    def test_subject_and_attachment_filename_are_scanned(self):
        message = EmailMessage()
        message["From"] = "legal@example.com"
        message["To"] = "review@example.com"
        message["Subject"] = "Patient SSN 321-54-9876"
        message.set_content("No additional body findings.")
        message.add_attachment(
            "reference only",
            filename="passport-12345678.txt",
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.eml"
            file_path.write_bytes(message.as_bytes())

            result = process_single_eml(file_path)
            detected_types = {match.pii_type for match in result.pii_matches}

        self.assertIn("SSN", detected_types)
        self.assertIn("PASSPORT", detected_types)

    def test_full_name_label_is_case_insensitive(self):
        text = "Full Name: Henry Robinson\nDOB: 11/08/1951\n"

        matches = scan_text(text, "record.txt")
        normalized_values = {match.pii_type: match.normalized_value for match in matches}

        self.assertIn("FULL_NAME", normalized_values)
        self.assertEqual("HENRY ROBINSON", normalized_values["FULL_NAME"])

    def test_service_role_email_is_suppressed_even_inside_record_context(self):
        text = (
            "Record owner: Grace Kim\n"
            "Support contact: support@servicehub.com\n"
            "Personal Email: grace.kim@gmail.com\n"
        )

        matches = scan_text(text, "record.txt")
        emails = sorted(match.normalized_value for match in matches if match.pii_type == "EMAIL")

        self.assertEqual(["grace.kim@gmail.com"], emails)

    def test_generic_coordinator_email_is_suppressed_even_when_personal_email_exists_elsewhere(self):
        text = (
            "Assigned coordinator: benefits@company.com\n"
            "Record owner: James Nguyen\n"
            "Personal Email: james.nguyen2@yahoo.com\n"
        )

        matches = scan_text(text, "record.txt")
        emails = sorted(match.normalized_value for match in matches if match.pii_type == "EMAIL")

        self.assertEqual(["james.nguyen2@yahoo.com"], emails)

    def test_forwarded_operational_header_email_is_not_treated_as_pii(self):
        text = (
            "Operations review team,\n"
            "No customer or employee record should be present in the quoted content.\n\n"
            "----- Forwarded message -----\n"
            "From: ops@healthsystem.org\n"
            "Support inbox: helpdesk@company.com\n"
            "Review date: 2026-02-11\n"
        )

        matches = scan_text(text, "ops_thread.txt")
        self.assertFalse(any(match.pii_type == "EMAIL" for match in matches))

    def test_login_ip_is_detected_when_personal_context_exists(self):
        text = "Customer portal login IP Address: 198.51.100.24\n"

        matches = scan_text(text, "record.txt")
        self.assertTrue(any(match.pii_type == "IPV4" for match in matches))

    def test_login_ip_is_not_suppressed_by_nearby_ticket_context(self):
        text = (
            "Reference ticket: FX987654321\n"
            "Questions line: (866) 555-0159\n"
            "Login IP Address: 69.141.155.217\n"
            "Personal Email: david.young5@outlook.com\n"
        )

        matches = scan_text(text, "record.txt")
        ip_values = [match.normalized_value for match in matches if match.pii_type == "IPV4"]

        self.assertEqual(["69141155217"], ip_values)

    def test_iban_detection_accepts_variable_group_lengths(self):
        text = "IBAN: FR14 2004 1010 0505 0001 3M02 606\n"

        matches = scan_text(text, "record.txt")
        normalized_values = {match.pii_type: match.normalized_value for match in matches}

        self.assertEqual("FR1420041010050500013M02606", normalized_values["IBAN"])

    def test_drivers_license_header_with_underscore_is_detected(self):
        text = "drivers_license_number,OR3733918\n"

        matches = scan_text(text, "record.csv")
        normalized_values = {match.pii_type: match.normalized_value for match in matches}

        self.assertEqual("OR3733918", normalized_values["DRIVERS_LICENSE"])

    def test_csv_full_name_field_is_detected(self):
        text = "Field,Value\nFull Name,Ethan Carter\nDOB,07/07/1957\n"

        matches = scan_text(text, "record.csv")
        names = [match.normalized_value for match in matches if match.pii_type == "FULL_NAME"]

        self.assertEqual(["ETHAN CARTER"], names)

    def test_xlsx_key_value_sheet_is_rendered_as_field_value_lines(self):
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "Record"
        worksheet.append(["field", "value"])
        worksheet.append(["Full Name", "Lucas Brooks"])
        worksheet.append(["DOB", "02/24/1957"])
        worksheet.append(["Driver's License Number", "AZ7756997"])
        worksheet.append(["Mobile Phone", "(939) 555-3355"])

        buffer = BytesIO()
        workbook.save(buffer)

        extracted = xlsx_extractor.extract(buffer.getvalue())
        matches = scan_text(extracted, "record.xlsx")
        normalized_values = {match.pii_type: match.normalized_value for match in matches}

        self.assertEqual("LUCAS BROOKS", normalized_values["FULL_NAME"])
        self.assertEqual("1957-02-24", normalized_values["DOB"])
        self.assertEqual("AZ7756997", normalized_values["DRIVERS_LICENSE"])
        self.assertEqual("9395553355", normalized_values["PHONE"])

    def test_xlsx_tabular_dependents_sheet_detects_multiple_full_names(self):
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "Dependents"
        worksheet.append(["role", "full name", "dob", "relationship", "ssn", "mrn", "phone"])
        worksheet.append(["employee", "Daniel Singh", "08/09/1975", "self", "540-49-8601", "243006", "(638) 555-4612"])
        worksheet.append(["dependent", "David Young", "12/14/1990", "child", "542-92-6955", "169024", "(257) 555-4937"])

        buffer = BytesIO()
        workbook.save(buffer)

        extracted = xlsx_extractor.extract(buffer.getvalue())
        matches = scan_text(extracted, "dependents.xlsx")
        names = sorted(match.normalized_value for match in matches if match.pii_type == "FULL_NAME")

        self.assertEqual(["DANIEL SINGH", "DAVID YOUNG"], names)

    def test_docx_table_rows_are_rendered_as_labeled_fields(self):
        document = Document()
        document.add_paragraph("Household Coverage")
        table = document.add_table(rows=3, cols=8)
        headers = ["role", "full name", "relationship", "dob", "ssn", "mrn", "phone", "personal email"]
        values = [
            ["primary", "William Taylor", "self", "12/28/1963", "212-97-1577", "704889", "(778) 555-5080", "william.taylor1@protonmail.com"],
            ["co-applicant", "William Taylor", "spouse", "04/11/1975", "249-64-5548", "112132", "(713) 555-1787", "william.taylor4@protonmail.com"],
        ]
        for row_idx, row_values in enumerate([headers, *values]):
            for col_idx, value in enumerate(row_values):
                table.cell(row_idx, col_idx).text = value

        buffer = BytesIO()
        document.save(buffer)

        extracted = extract_text_from_attachment(
            Attachment(
                filename="household.docx",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=buffer.getvalue(),
                source_eml="sample.eml",
            )
        )
        matches = scan_text(f"Attachment filename: household.docx\n{extracted}", "sample.eml > household.docx")
        detected_types = {match.pii_type for match in matches}
        names = sorted(match.normalized_value for match in matches if match.pii_type == "FULL_NAME")
        mrns = sorted(match.normalized_value for match in matches if match.pii_type == "MRN")

        self.assertTrue({"FULL_NAME", "MRN", "DOB", "SSN", "PHONE", "EMAIL"} <= detected_types)
        self.assertEqual(["WILLIAM TAYLOR", "WILLIAM TAYLOR"], names)
        self.assertEqual(["112132", "704889"], mrns)

    def test_csv_table_rows_are_rendered_as_labeled_fields(self):
        csv_text = (
            "role,full name,relationship,dob,ssn,mrn,phone,personal email\n"
            "primary,Sophia Miller,self,01/05/1960,302-79-5850,320207,(567) 555-3454,sophia.miller7@email.com\n"
            "co-applicant,Sophia Miller,spouse,08/04/1972,575-48-6251,976660,(241) 555-9319,sophia.miller3@email.com\n"
        )

        extracted = extract_text_from_attachment(
            Attachment(
                filename="household.csv",
                mime_type="text/csv",
                data=csv_text.encode("utf-8"),
                source_eml="sample.eml",
            )
        )
        matches = scan_text(f"Attachment filename: household.csv\n{extracted}", "sample.eml > household.csv")
        names = [match for match in matches if match.pii_type == "FULL_NAME"]
        mrns = sorted(match.normalized_value for match in matches if match.pii_type == "MRN")

        self.assertEqual(2, len(names))
        self.assertEqual(["320207", "976660"], mrns)

    def test_zip_embedded_office_record_uses_extension_aware_extraction(self):
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "Record"
        worksheet.append(["field", "value"])
        worksheet.append(["Full Name", "Benjamin Washington"])
        worksheet.append(["DOB", "04/18/1993"])
        worksheet.append(["SSN", "590-10-6481"])
        worksheet.append(["Address", "4911 Hillcrest Lane, Tampa, FL 33606"])

        record_bytes = BytesIO()
        workbook.save(record_bytes)

        archive_bytes = BytesIO()
        with zipfile.ZipFile(archive_bytes, "w") as archive:
            archive.writestr("record_103.xlsx", record_bytes.getvalue())

        extracted = extract_text_from_attachment(
            Attachment(
                filename="review_packet_103.zip",
                mime_type="application/zip",
                data=archive_bytes.getvalue(),
                source_eml="sample.eml",
            )
        )
        matches = scan_text(
            f"Attachment filename: review_packet_103.zip\n{extracted}",
            "sample.eml > review_packet_103.zip",
        )
        detected_types = {match.pii_type for match in matches}

        self.assertTrue({"FULL_NAME", "DOB", "SSN", "ADDRESS"} <= detected_types)


if __name__ == "__main__":
    unittest.main()
